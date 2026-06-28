"""Markdown → Word(.docx)：用 python-docx 把报告正文转成可下载的 Word 文档。

支持：标题(#~######)、段落、有序/无序列表、引用块(>)、表格(| |)、分隔线、
代码块(```)，以及行内 **加粗** / `代码` / [文字](链接)。
设计为「尽量保留结构、解析失败降级为纯段落」，不追求像素级还原。
"""
import io
import re

from app.core.logging import get_logger

logger = get_logger(__name__)

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_ULIST_RE = re.compile(r"^[-*]\s+(.*)$")
_OLIST_RE = re.compile(r"^\d+\.\s+(.*)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
# 行内：**加粗** / `代码` / [文字](链接)
_INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
_LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")


def _is_table_sep(line: str) -> bool:
    s = line.strip()
    return bool(s) and set(s) <= set("|:- ") and "-" in s and "|" in s


def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _add_inline(paragraph, text: str) -> None:
    """把行内 markdown 标记转成 docx run（加粗/代码/链接降级为文字）。"""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
        else:
            lm = _LINK_RE.match(token)
            if lm:
                # 链接降级为「文字」纯文本（docx 超链接较繁琐，正文以可读为先）
                paragraph.add_run(lm.group(1))
            else:
                paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx_bytes(title: str, markdown: str) -> bytes:
    """把 Markdown 转成 .docx 字节流。解析异常时降级为纯文本段落。"""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    if title:
        doc.add_heading(title, level=0)

    lines = (markdown or "").splitlines()
    i = 0
    n = len(lines)
    in_code = False
    code_buf: list[str] = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 代码块围栏
        if stripped.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                run = para.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # 表格：当前行含 | 且下一行是分隔行
        if "|" in stripped and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _split_row(stripped)
            rows: list[list[str]] = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                rows.append(_split_row(lines[j]))
                j += 1
            try:
                table = doc.add_table(rows=1, cols=len(header))
                table.style = "Light Grid Accent 1"
                for c, text in enumerate(header):
                    table.rows[0].cells[c].text = text
                for r in rows:
                    cells = table.add_row().cells
                    for c in range(len(header)):
                        cells[c].text = r[c] if c < len(r) else ""
            except Exception as e:  # noqa: BLE001
                logger.warning("Word 表格转换失败，降级文本: %s", e)
                doc.add_paragraph(stripped)
            i = j
            continue

        # 标题
        hm = _HEADING_RE.match(stripped)
        if hm:
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2).strip(), level=level)
            i += 1
            continue

        # 分隔线
        if _HR_RE.match(stripped):
            doc.add_paragraph().add_run("─" * 20)
            i += 1
            continue

        # 引用
        if stripped.startswith(">"):
            para = doc.add_paragraph(style="Intense Quote")
            _add_inline(para, stripped.lstrip("> ").strip())
            i += 1
            continue

        # 列表
        um = _ULIST_RE.match(stripped)
        if um:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, um.group(1).strip())
            i += 1
            continue
        om = _OLIST_RE.match(stripped)
        if om:
            para = doc.add_paragraph(style="List Number")
            _add_inline(para, om.group(1).strip())
            i += 1
            continue

        # 普通段落
        para = doc.add_paragraph()
        _add_inline(para, stripped)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
